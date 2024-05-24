import hashlib as hs
import random
import sympy
import tkinter as tk
from tkinter import font as tkfont
from tkinter import messagebox
from tkinter import filedialog
import docx
from docx import Document
from tkhtmlview import HTMLLabel
from tkinter import Tk, Text, Scrollbar, END
from tkinter.scrolledtext import ScrolledText
from docx.shared import RGBColor

# CÁC THUẬT TOÁN SỬ DỤNG - THUẬT TOÁN DSA
# khai báo biến toàn cục
global l, q, x, y, p, g, r, s
global file_path
file_path = ""
global check


# BĂM DỮ LIỆU
def bamDuLieu(data):
    text = hs.sha1(data.encode()).hexdigest()  # Ham bam mat ma
    hexa_to_bytes = bytes.fromhex(text)  # Chuyen ham bam sang dang byte
    hexa_to_int = int.from_bytes(hexa_to_bytes, byteorder='big')  # Ham doi hexa sang integer
    return hexa_to_int


# THUẬT TOÁN BÌNH PHƯƠNG VÀ NHÂN
def binhPhuongvaNhan(x, n, p):
    result = 1  # Khởi tạo tích p = 1
    binary_n = bin(n)[2:]  # Lấy giá trị nhị phân
    for bit in binary_n:
        result = (result ** 2) % p

        if bit == '1':  # Nếu bit là 1 thì nhân với x
            result = (result * x) % p
    return result


# THUẬT TOÁN ECULID MỞ RỘNG
def oClitMoRong(r0, r1):
    a, b = r0, r1
    sa, ta = 1, 0
    sb, tb = 0, 1
    if b == 1:
        return 0
    elif b < 1:
        return
    else:
        while (b > 1):  # Làm việc khi nào thương chưa bằng 1 - dư chưa bằng 0
            q = a // b
            r = a % b
            sr = sa - q * sb
            tr = ta - q * tb
            a = b
            sa, ta = sb, tb
            b = r
            sb, tb = sr, tr
            while (tr < 0):
                tr += r0
        return tr


# KỞI TAỌ THAM SỐ
def khoiTaoThamSo():
    global l, q, x, y, p, g
    l = random.randint(512, 1024)
    while (l % 64 != 0):
        l = random.randint(512, 1024)

    q = random.getrandbits(160)
    while sympy.isprime(q) != True:
        q = random.getrandbits(160)

    k = random.getrandbits(l - 160)
    p = k * q + 1

    while sympy.isprime(p) != True:
        k = random.getrandbits(l - 160)
        p = k * q + 1

    h = random.randint(2, p - 2)
    g = binhPhuongvaNhan(h, k, p)
    while g <= 1:
        h = random.randint(2, p - 2)
        g = binhPhuongvaNhan(h, k, p)

    x = random.randint(1, q - 1)  # Khoa rieng tu
    y = binhPhuongvaNhan(g, x, p)  # Khoa cong khai
    y = int(y)
    print("Cac gia tri ngau nhien:\n p={0},\n q={1},\n g={2},\n h={3}".format(p, q, g, h))
    print("Khoa bi mat x la: ", x)
    print("khoa cong khai y la: ", y)


# KÝ CHỮ KÝ ĐIỆN TỬ
def chuKyDienTu():
    global q, x, y, p, g, r, s
    khoiTaoThamSo()

    thongdiep = text_content.get("1.0", "end-1c")  # Thông điệp lấy từ nội dung text
    soLieuBam = bamDuLieu(thongdiep)  # Giá trị hàm băm được đổi ra cơ số 10

    k = random.randint(1, q - 1)
    print("Khoa random k la: ", k)
    r = binhPhuongvaNhan(g, k, p) % q
    s = ((oClitMoRong(q, k) * (soLieuBam + x * r)) % q) % q

    while (s == 0 or r == 0):  # Nếu r hoặc s =0 thì làm lại.
        k = random.randint(1, q - 1)
        r = binhPhuongvaNhan(g, k, p) % q
        s = ((oClitMoRong(q, k) * (soLieuBam + x * r)) % q) % q

    text_signature.delete("1.0", "end-1c")
    text_signature.insert('1.0', 'r= {0}\ns= {1}'.format(r, s))
    text_hash.delete("1.0", "end-1c")
    text_hash.insert("1.0", "H(M)= {0}".format(soLieuBam))
    print('Gia tri cap gia tri (r, s): ({0}, {1})'.format(int(r), int(s)))
    print(p, q, g)


# XÁC MINH CHỮ KÝ ĐIỆN TỬ
def kiemTraChuKy():
    global q, y, p, g, r, s
    # Băm lại nội dung ở đầu người nhận
    noidung = text_sent_content.get("1.0", "end-1c")
    gtrBamLai = bamDuLieu(noidung)
    text_hash_again.delete("1.0", "end-1c")
    text_hash_again.insert("1.0", "H(M)= {0}".format(gtrBamLai))

    if not (0 < r < q) or not (0 < s < q):
        return False
    else:
        w = oClitMoRong(q, s)
        u1 = (gtrBamLai * w) % q
        u1 = int(u1)
        u2 = (r * w) % q
        u2 = int(u2)
        v = ((binhPhuongvaNhan(g, u1, p) * binhPhuongvaNhan(y, u2, p)) % p) % q
        print("Cac gia tri p, q, r, s, g la: ", p, q, r, s, g)
        print('Gia tri r= {0},v= {1},solieubamlai= {2}'.format(int(r), int(v), gtrBamLai))
        return v == r


# ĐIỂU KHIỂN GIAO DIỆN NGƯỜI DÙNG
def taoChuKy():
    chuKyDienTu()
    # Hiển thị thông báo sinh chữ ký thành công
    messagebox.showinfo("Thông báo", "Sinh chữ ký thành công")


# CHUYỂN MÀU TỪ RGB SANG HEXA
def rgb_to_hex(rgb):
    return '#%02x%02x%02x' % (rgb[0], rgb[1], rgb[2])


# XÓA BỘ NHỚ CÓ SẴN
def xoaBoNho():
    global l, q, x, y, p, g, r, s
    global file_path
    file_path = ""
    l, q, x, y, p, g, r, s = 0, 0, 0, 0, 0, 0, 0, 0,
    messagebox.showinfo("Thông báo", "Xóa bộ nhớ thành công.")


def chonFileNoiDungGui():
    global file_path
    global content
    file_path = filedialog.askopenfilename(
        title='Chọn file',
        filetypes=[('Text Files', '*.txt'), ('Docx Files', '*.docx')]
    )
    try:
        if not file_path:
            messagebox.showwarning("Thông báo", "Bạn chưa chọn file !")
            return

        content = []
        if file_path.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as f:
                messagebox.showinfo("Thông báo", "Mở file thành công!")
                content = [(line, None) for line in f]
        elif file_path.endswith('.docx'):
            document = Document(file_path)
            for para in document.paragraphs:
                for run in para.runs:
                    if run.font.color and run.font.color.rgb:
                        color = run.font.color.rgb
                        content.append((run.text, rgb_to_hex((color[0], color[1], color[2]))))
                    else:
                        content.append((run.text, None))
            messagebox.showinfo("Thông báo", "Mở file thành công!")

        text_content.delete("1.0", "end-1c")
        for text, color in content:
            if color:
                text_content.insert(tk.END, text, color)
                text_content.tag_configure(color, foreground=color)
            else:
                text_content.insert(tk.END, text)
    except Exception as e:
        messagebox.showwarning("Thông báo", str(e))


def chonFileNoiDungNhan():
    global file_path
    global sentContent
    file_path = filedialog.askopenfilename(
        title='Chọn file',
        filetypes=[('Text Files', '*.txt'), ('Docx Files', '*.docx')]
    )
    try:
        if not file_path:
            messagebox.showwarning("Thông báo", "Bạn chưa chọn file !")
            return

        sentContent = []
        if file_path.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as f:
                messagebox.showinfo("Thông báo", "Mở file thành công!")
                sentContent = [(line, None) for line in f]
        elif file_path.endswith('.docx'):
            document = Document(file_path)
            for para in document.paragraphs:
                for run in para.runs:
                    if run.font.color and run.font.color.rgb:
                        color = run.font.color.rgb
                        sentContent.append((run.text, rgb_to_hex((color[0], color[1], color[2]))))
                    else:
                        sentContent.append((run.text, None))
            messagebox.showinfo("Thông báo", "Mở file thành công!")

        text_sent_content.delete("1.0", "end-1c")
        for text, color in sentContent:
            if color:
                text_sent_content.insert(tk.END, text, color)
                text_sent_content.tag_configure(color, foreground=color)
            else:
                text_sent_content.insert(tk.END, text)
    except Exception as e:
        messagebox.showwarning("Thông báo", "Có gì đó sai sai !!!")


def chonFileChuKy():
    global sentSignature
    file_path = filedialog.askopenfilename(
        title='Chọn file',
        filetypes=[('Text Files', '*.txt'), ('Docx Files', '*.docx')]
    )
    try:
        if not file_path:
            messagebox.showwarning("Thông báo", "Bạn chưa chọn file !")
            return

        with open(file_path, 'r') as f:
            messagebox.showinfo("Thông báo", "Mở file thành công !")
            sentSignature = f.read()
            sentSignature = sentSignature.strip("{}")
            text_sent_content_signature.delete("1.0", "end-1c")
            text_sent_content_signature.insert("1.0", sentSignature)
    except Exception as e:
        messagebox.showwarning("Thông báo", "Bạn chưa chọn file")


def luuChuKy():
    saved_filePath = filedialog.asksaveasfilename(
        title="Chọn nơi lưu",
        defaultextension=".txt"
    )
    if saved_filePath:  # Kiểm tra nếu người dùng không hủy hộp thoại
        try:
            with open(saved_filePath, 'w', encoding='utf-8') as f:
                f.write(text_signature.get("1.0", "end-1c"))
            messagebox.showinfo("Thông báo", "Lưu chữ ký thành công!")
        except Exception as e:
            print(e)
            messagebox.showerror("Thông báo", "Có gì đó sai sai!!!")


def guiNoiDung():
    # Gửi nội dung
    global file_path
    global sentContent
    text_sent_content.delete("1.0", "end-1c")
    if file_path == "" or file_path.endswith(".txt"):
        text_sent_content.insert("1.0", text_content.get("1.0", "end-1c"))
    else:
        sentContent = []
        document = Document(file_path)
        for para in document.paragraphs:
            for run in para.runs:
                if run.font.color and run.font.color.rgb:
                    color = run.font.color.rgb
                    sentContent.append((run.text, rgb_to_hex((color[0], color[1], color[2]))))
                else:
                    sentContent.append((run.text, None))
        for text, color in sentContent:
            if color:
                text_sent_content.insert(tk.END, text, color)
                text_sent_content.tag_configure(color, foreground=color)
            else:
                text_sent_content.insert(tk.END, text)

    # Gửi bản chữ ký
    text_sent_content_signature.delete("1.0", "end-1c")
    text_sent_content_signature.insert("1.0", text_signature.get("1.0", "end-1c"))

    messagebox.showinfo("Thông báo", "Gửi nội dung thành công")


def xacThucNoiDung():
    global check
    check = False
    if (text_signature.get("1.0", "end-1c") == text_sent_content_signature.get("1.0", "end-1c")):
        check = True

    if kiemTraChuKy() == True and check == True:
        messagebox.showinfo('Thông báo', 'Xác thực thành công')
        text_verify.delete("1.0", "end-1c")
        text_verify.insert("1.0", "Chữ ký trùng khớp" + "\nVăn bản nguyên vẹn")
    elif kiemTraChuKy() == False and check == True:
        messagebox.showerror('Thông báo', 'Xác thực chữ ký thất bại')
        text_verify.delete("1.0", "end-1c")
        text_verify.insert("1.0", "Chữ ký trùng khớp" + "\nVăn bản đã bị thay đổi")
    elif kiemTraChuKy() == True and check == False:
        messagebox.showerror('Thông báo', 'Xác thực chữ ký thất bại')
        text_verify.delete("1.0", "end-1c")
        text_verify.insert("1.0", "Chữ ký không trùng khớp" + "\nVăn bản nguyên vẹn")
    else:
        messagebox.showerror('Thông báo', 'Xác thực chữ ký thất bại')
        text_verify.delete("1.0", "end-1c")
        text_verify.insert("1.0", "Chữ ký không trùng khớp" + "\nVăn bản đã bị thay đổi")


# THỰC THI GIAO TIỆN NGƯỜI DÙNG
window = tk.Tk()
window.title("CHỮ KÝ SỐ")
window.geometry("1050x550+180+50")
window.iconbitmap("D:\Learning IT\Python\BTL\calculator.ico")

main_frame = tk.Frame(window)
main_frame.pack(fill=tk.BOTH, expand=True)

bg_left = 'lightblue'
left_frame = tk.Frame(main_frame, bg=bg_left, width=400, height=400)
left_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

bg_right = 'lightgreen'
right_frame = tk.Frame(main_frame, bg=bg_right, width=400, height=400)
right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)

bigger_font = tkfont.Font(size=10)

# XÂY DỰNG LAYOUT BÊN TRÁI
label_content = tk.Label(left_frame, text='NỘI DUNG: ', bg=bg_left, fg='red')
label_content.grid(column=0, row=0, padx=5, pady=20)
text_content = tk.Text(left_frame, font=bigger_font, height=8, width=45)
text_content.grid(column=1, row=0, padx=5, pady=20)
btn_choose_file = tk.Button(left_frame, text='Chọn File', bg=bg_right, fg='red', bd=5, command=chonFileNoiDungGui)
btn_choose_file.grid(column=2, row=0, padx=5, pady=20)
btn_encode = tk.Button(left_frame, text='Tạo Chữ Ký Số', bg=bg_right, fg='red', bd=5, command=taoChuKy)
btn_encode.grid(column=1, row=2, columnspan=3)

label_hash = tk.Label(left_frame, text='HÀM BĂM: ', bg=bg_left, fg='red')
label_hash.grid(column=0, row=3, padx=5, pady=20)
text_hash = tk.Text(left_frame, font=bigger_font, height=3, width=45)
text_hash.grid(column=1, row=3, padx=5, pady=20)

label_signature = tk.Label(left_frame, text='CHỮ KÝ: ', fg='red', bg=bg_left)
label_signature.grid(column=0, row=4, padx=5, pady=20)
text_signature = tk.Text(left_frame, font=bigger_font, height=4, width=45)
text_signature.grid(column=1, row=4, padx=5, pady=20)
btn_send = tk.Button(left_frame, text='Gửi nội dung', fg='red', bg=bg_right, bd=5, command=guiNoiDung)
btn_send.grid(column=2, row=4, padx=5, pady=20)

btn_saveSignature = tk.Button(left_frame, text='Lưu file chữ ký', bd=5, fg='red', bg=bg_right, command=luuChuKy)
btn_saveSignature.grid(column=2, row=5, columnspan=2, padx=5, pady=30)

btn_resetData = tk.Button(left_frame, text="Xóa bộ nhớ", bd=5, fg='red', bg=bg_right, command=xoaBoNho)
btn_resetData.grid(column=0, row=5, columnspan=2, padx=5, pady=30)

# XÂY DỰNG LAYOUT BÊN PHẢI
label_sent_content = tk.Label(right_frame, text='NỘI DUNG:', fg='red', bg=bg_right)
label_sent_content.grid(row=0, column=0, padx=5, pady=20)

text_sent_content = tk.Text(right_frame, font=bigger_font, height=8, width=45)
text_sent_content.grid(column=1, row=0, padx=5, pady=20)

btn_choose_file_content = tk.Button(right_frame, text='Chọn File', bg=bg_left, fg='red', bd=5,
                                    command=chonFileNoiDungNhan)
btn_choose_file_content.grid(column=2, row=0, padx=5, pady=20)

label_sent_signature = tk.Label(right_frame, text='CHỮ KÝ: ', fg='red', bg=bg_right)
label_sent_signature.grid(column=0, row=1, padx=5, pady=20)

text_sent_content_signature = tk.Text(right_frame, font=bigger_font, height=4, width=45)
text_sent_content_signature.grid(column=1, row=1, padx=5, pady=20)

btn_choose_file_signature = tk.Button(right_frame, text='Chọn File', bg=bg_left, fg='red', bd=5, command=chonFileChuKy)
btn_choose_file_signature.grid(column=2, row=1, padx=5, pady=20)

btn_verify = tk.Button(right_frame, text='Kiểm tra chữ ký', fg='red', bg=bg_left, bd=5, command=xacThucNoiDung)
btn_verify.grid(column=1, row=2, columnspan=3, padx=5, pady=2)

label_hash_again = tk.Label(right_frame, text='HÀM BĂM: ', bg=bg_right, fg='red')
label_hash_again.grid(column=0, row=3, padx=5, pady=20)
text_hash_again = tk.Text(right_frame, font=bigger_font, height=3, width=45)
text_hash_again.grid(column=1, row=3, padx=5, pady=20)

label_verify = tk.Label(right_frame, text='THÔNG BÁO:', bg=bg_right, fg='red')
label_verify.grid(column=0, row=4, padx=5, pady=20)

text_verify = tk.Text(right_frame, font=bigger_font, height=4, width=45)
text_verify.grid(column=1, row=4, pady=20)

window.mainloop()
